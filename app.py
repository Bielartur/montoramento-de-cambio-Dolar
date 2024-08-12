from selenium.webdriver import Chrome
from selenium.webdriver import ChromeOptions
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import datetime
from time import sleep
import os
import sys

# Salvar nome/link do site escolhido para cotação do Dólar para Real
SITE = 'https://www.xe.com/pt/currencyconverter/convert/?Amount=1&From=USD&To=BRL'

# Configurar Driver
base_path = os.path.dirname(sys.argv[0])
chromedriver_path = os.path.join(base_path, "Driver", "chromedriver.exe")
options = ChromeOptions()
options.add_argument('--icognito')
options.add_argument('--start-maximized')
driver = Chrome(options=options)

def entrar_no_site():
    # Entrar no site
    driver.get(SITE)
    wait = WebDriverWait(driver, 10)
    sleep(1)


def extrair_cotacao():
    # Coletar o valor do Dólar para o dia atual
    cotacao = driver.find_element(By.XPATH, '/html/body/div[1]/div/div[4]/div[2]/section/div[2]/div/main/div/div[2]/div[1]/div/p[2]')
    cotacao = cotacao.text
    cotacao = cotacao[:4]
    return cotacao


def tirar_print():
    # Simular o "apertar F11" para um print mais claro
    actions = ActionChains(driver)
    actions.send_keys(Keys.F11).perform()
    sleep(1)
    # Salvar print do site
    driver.save_screenshot('Cotação Dolar.png')


def fechar_driver():
    driver.quit()


def reconhecer_data_atual():
    # Salvar a data a qual foi acessado o site
    data_atual = datetime.date.today()
    data_atual_formatada = data_atual.strftime('%d/%m/%Y')

    data_atual = str(data_atual_formatada)
    data_atual = data_atual.replace('-', '/')

    return data_atual


def escrever_documento():
    doc = Document()

    # Centralizar deixar em negrito o título
    titulo = doc.add_paragraph()
    run = titulo.add_run(f'Cotação Atual do Dólar – {cotacao} ({data_atual})')
    run.bold = True
    run.font.size = Pt(18)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adicionar texto do valor da cotação
    texto = doc.add_paragraph()
    run2 = texto.add_run(f'O dólar está no valor de {cotacao}, na data {data_atual}.')
    run2.font.size = Pt(12)

    # Adicionar texto informano o site
    texto_site = doc.add_paragraph()
    run3 = texto_site.add_run(f'Valor cotado no site: ')
    run3.font.size = Pt(12)
    adicionar_hiperlink(texto_site, SITE, SITE)

    # Adicionando a imagem
    texto_descreve_imagem = doc.add_paragraph()
    run4 = texto_descreve_imagem.add_run('Print da cotação atual')
    run4.font.size = Pt(12)
    doc.add_picture('Cotação Dolar.png', width=Inches(6.0))

    # Adicionar texto informando a autoria
    texto_autoria = doc.add_paragraph()
    run5 = texto_autoria.add_run('Cotação feita por - Gabriel Artur S S Canto')
    run5.font.size = Pt(12)

    nome_documento = 'Cotação Dólar.docx'
    doc.save(nome_documento)
    print('Documento escrito com sucesso!')

    return nome_documento


def adicionar_hiperlink(paragrafo, texto, url):
    part = paragrafo.part
    r_id = part.relate_to(url, 'hyperlink', is_external=True)

    # Cria o elemento <w:hyperlink>
    hiperlink = OxmlElement('w:hyperlink')
    hiperlink.set(qn('r:id'), r_id)

    # Cria o elemento <w:r> com o texto do hyperlink
    nova_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    nova_run.append(rPr)
    text_run = OxmlElement('w:t')
    text_run.text = texto
    nova_run.append(text_run)

    # Anexa o <w:r> ao <w:hyperlink>
    hiperlink.append(nova_run)

    # Anexa o <w:hyperlink> ao parágrafo
    paragrafo._p.append(hiperlink)


def converter_documento_para_pdf(caminho_documento):
    # Tornar esses dados em um relatório PDF
    pass


entrar_no_site()
cotacao = extrair_cotacao()
tirar_print()
fechar_driver()
data_atual = reconhecer_data_atual()
caminho_documento = escrever_documento()
converter_documento_para_pdf(caminho_documento)
sleep(3)
